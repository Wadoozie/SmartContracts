"""
Generate Wadoozie DAO — Smart Contract Reference (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                      "Wadoozie_DAO_Contract_Reference.docx")

doc = Document()

# ── Global style tweaks ──────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def shade_table(table):
    """Shade header row and apply borders."""
    for cell in table.rows[0].cells:
        shading = cell._element.get_or_add_tcPr()
        sh_elem = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A1A2E",
            qn("w:val"): "clear",
        })
        shading.append(sh_elem)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def add_table(headers, rows):
    """Add a formatted table with shaded header."""
    t = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    shade_table(t)
    doc.add_paragraph()  # spacer
    return t


def bold_para(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

title = doc.add_heading("Wadoozie DAO — Smart Contract Reference", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("Contract Methods & Deployment Guide")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Overview", level=1)

doc.add_paragraph(
    "Wadoozie DAO is a fully on-chain governance system built on Ethereum. "
    "Token holders vote on proposals that control the DAO treasury and can "
    "change governance parameters — no single person has admin power once "
    "deployment is complete."
)

doc.add_heading("The Three Contracts", level=2)

add_table(
    ["Contract", "Role"],
    [
        ["Wadoozie (WADZ)", "The governance token. Holding it gives you voting power (after delegation)."],
        ["Headquarters (Governor)", "The governance hub. Proposals are created, voted on, queued, and executed here."],
        ["WadoozieTimelock (Treasury)", "The time-delayed treasury. Approved proposals wait here before execution, giving the community a window to react."],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. WADOOZIE TOKEN
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Wadoozie Token (WADZ)", level=1)

doc.add_paragraph(
    "Wadoozie is a standard ERC-20 token with a fixed supply of 1 billion "
    "tokens. It cannot be minted, burned, or paused after deployment — the "
    "supply is permanently locked. It supports ERC-20 Permit (gasless "
    "approvals) and ERC-20 Votes (on-chain voting power tracking)."
)

doc.add_heading("Constructor Parameters", level=2)

add_table(
    ["Parameter", "Type", "Description"],
    [
        ["initialHolder", "address", "Wallet that receives the entire 1 billion token supply at deployment."],
    ],
)

doc.add_heading("Read Methods (free, no gas)", level=2)

add_table(
    ["Method", "Returns", "What It Does"],
    [
        ["name()", "string", "Returns the token name (\"Wadoozie\" on testnet, \"Wadoozie\" on mainnet)."],
        ["symbol()", "string", "Returns the token symbol (\"WADZ\" on testnet, \"WADZ\" on mainnet)."],
        ["decimals()", "uint8", "Returns 18 (standard ERC-20 precision)."],
        ["totalSupply()", "uint256", "Returns the total token supply (1,000,000,000 * 10^18)."],
        ["TOTAL_SUPPLY()", "uint256", "Same as totalSupply — a public constant for convenience."],
        ["balanceOf(account)", "uint256", "Returns how many tokens the given address holds."],
        ["allowance(owner, spender)", "uint256", "Returns how many tokens a spender is approved to transfer on behalf of owner."],
        ["delegates(account)", "address", "Returns who the account has delegated their voting power to."],
        ["getVotes(account)", "uint256", "Returns current voting power of an address (includes delegated power)."],
        ["getPastVotes(account, timepoint)", "uint256", "Returns voting power at a specific past block number."],
        ["getPastTotalSupply(timepoint)", "uint256", "Returns the total voting supply at a past block."],
        ["nonces(owner)", "uint256", "Returns the current permit nonce for gasless approvals (EIP-2612)."],
        ["clock()", "uint48", "Returns the current clock value used by the voting system (block number)."],
        ["CLOCK_MODE()", "string", "Returns \"mode=blocknumber&from=default\" — describes how the clock works."],
    ],
)

doc.add_heading("Write Methods (require gas)", level=2)

add_table(
    ["Method", "Parameters", "What It Does"],
    [
        ["transfer(to, amount)", "to: address, amount: uint256", "Send tokens from your wallet to another address."],
        ["approve(spender, amount)", "spender: address, amount: uint256", "Allow another address to spend up to amount of your tokens."],
        ["transferFrom(from, to, amount)", "from: address, to: address, amount: uint256", "Transfer tokens on behalf of another address (requires prior approval)."],
        ["delegate(delegatee)", "delegatee: address", "Assign your voting power to another address (or yourself). Required before you can vote."],
        ["delegateBySig(delegatee, nonce, expiry, v, r, s)", "See EIP-712", "Delegate voting power using an off-chain signature (gasless for the token holder)."],
        ["permit(owner, spender, value, deadline, v, r, s)", "See EIP-2612", "Approve a spender using an off-chain signature instead of a transaction."],
    ],
)

doc.add_heading("Key Concept: Delegation", level=2)

doc.add_paragraph(
    "Holding WADZ tokens alone does NOT give you voting power. You must "
    "delegate your tokens — either to yourself or to someone else — before "
    "they count as votes. If you want to vote yourself, call "
    "delegate(yourOwnAddress). Delegation is a one-time action per address; "
    "it stays in effect until you change it."
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. HEADQUARTERS (GOVERNOR)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Headquarters (Governor)", level=1)

doc.add_paragraph(
    "Headquarters is the governance contract — the decision-making engine of "
    "the DAO. Anyone with enough delegated voting power can create a "
    "proposal. Token holders vote For, Against, or Abstain. If a proposal "
    "passes and reaches quorum, it is queued in the Timelock and executed "
    "after a delay."
)

doc.add_heading("Constructor Parameters", level=2)

add_table(
    ["Parameter", "Type", "Description", "Recommended Value"],
    [
        ["_token", "address", "Address of the deployed Wadoozie token.", "(deployed address)"],
        ["_timelock", "address", "Address of the deployed WadoozieTimelock.", "(deployed address)"],
        ["_votingDelay", "uint48", "Number of blocks between proposal creation and voting start.", "7,200 (~1 day)"],
        ["_votingPeriod", "uint32", "Number of blocks the voting window stays open.", "50,400 (~1 week)"],
        ["_proposalThreshold", "uint256", "Minimum delegated voting power required to create a proposal (in token units with 18 decimals).", "1,000 WADZ"],
        ["_quorumPercent", "uint256", "Percentage of total supply that must vote for a proposal to be valid.", "4 (= 4%)"],
        ["_guardian", "address", "Emergency address that can cancel any active proposal. Must be non-zero at deploy. Use a multisig.", "(multisig address)"],
    ],
)

doc.add_heading("Governance Lifecycle", level=2)

doc.add_paragraph(
    "Every proposal follows a fixed lifecycle. Each step must complete "
    "before the next one can begin."
)

add_table(
    ["Step", "Action", "Who Can Do It", "What Happens"],
    [
        ["1. Propose", "propose(targets, values, calldatas, description)", "Any address with >= proposalThreshold delegated votes", "A new proposal is created. Voting begins after the voting delay."],
        ["2. Vote", "castVote(proposalId, support)  or  castVoteWithReason(...)", "Any address with delegated voting power", "Token holders vote For (1), Against (0), or Abstain (2) during the voting period."],
        ["3. Queue", "queue(targets, values, calldatas, descriptionHash)", "Anyone (if proposal succeeded)", "The passed proposal is sent to the Timelock. A countdown begins."],
        ["4. Execute", "execute(targets, values, calldatas, descriptionHash)", "Anyone (after timelock delay)", "The on-chain actions in the proposal are carried out (transfers, parameter changes, etc.)."],
    ],
)

doc.add_heading("Read Methods (free, no gas)", level=2)

add_table(
    ["Method", "Returns", "What It Does"],
    [
        ["name()", "string", "Returns \"Headquarters\"."],
        ["version()", "string", "Returns the Governor contract version."],
        ["token()", "address", "Returns the address of the governance token."],
        ["timelock()", "address", "Returns the address of the Timelock."],
        ["votingDelay()", "uint256", "Blocks between proposal creation and vote start."],
        ["votingPeriod()", "uint256", "Blocks the voting window stays open."],
        ["proposalThreshold()", "uint256", "Minimum voting power to create a proposal."],
        ["quorum(blockNumber)", "uint256", "Minimum total votes needed for a proposal at a given block."],
        ["quorumNumerator()", "uint256", "The quorum percentage numerator (e.g. 4)."],
        ["quorumDenominator()", "uint256", "The quorum percentage denominator (100)."],
        ["state(proposalId)", "ProposalState", "Current state of a proposal (Pending, Active, Canceled, Defeated, Succeeded, Queued, Expired, Executed)."],
        ["proposalNeedsQueuing(proposalId)", "bool", "Whether a proposal must be queued before execution (always true for timelock-governed)."],
        ["proposalSnapshot(proposalId)", "uint256", "Block number at which votes are counted."],
        ["proposalDeadline(proposalId)", "uint256", "Block number when voting ends."],
        ["proposalEta(proposalId)", "uint256", "Timestamp when a queued proposal can be executed."],
        ["proposalProposer(proposalId)", "address", "Address that created the proposal."],
        ["hasVoted(proposalId, account)", "bool", "Whether an address has already voted on a proposal."],
        ["proposalVotes(proposalId)", "againstVotes, forVotes, abstainVotes", "Current vote tallies for a proposal."],
        ["proposalGuardian()", "address", "Returns the current proposal guardian address."],
        ["getVotes(account, timepoint)", "uint256", "Voting power of an address at a specific block."],
        ["clock()", "uint48", "Current clock value (block number)."],
        ["CLOCK_MODE()", "string", "Clock mode description."],
        ["COUNTING_MODE()", "string", "Describes how votes are counted (\"support=bravo&quorum=for,abstain\")."],
    ],
)

doc.add_heading("Write Methods (require gas)", level=2)

add_table(
    ["Method", "Parameters", "What It Does"],
    [
        ["propose(targets, values, calldatas, description)", "targets: address[], values: uint256[], calldatas: bytes[], description: string", "Create a new governance proposal. Caller must meet the proposal threshold."],
        ["castVote(proposalId, support)", "proposalId: uint256, support: uint8 (0=Against, 1=For, 2=Abstain)", "Cast your vote on an active proposal."],
        ["castVoteWithReason(proposalId, support, reason)", "proposalId: uint256, support: uint8, reason: string", "Cast your vote with an on-chain reason string."],
        ["castVoteWithReasonAndParams(...)", "proposalId, support, reason, params", "Cast your vote with reason and additional parameters."],
        ["castVoteBySig(proposalId, support, voter, signature)", "See EIP-712", "Vote using an off-chain signature (gasless for the voter)."],
        ["queue(targets, values, calldatas, descriptionHash)", "Same as propose, but descriptionHash is bytes32", "Queue a succeeded proposal into the Timelock."],
        ["execute(targets, values, calldatas, descriptionHash)", "Same as queue", "Execute a queued proposal after the timelock delay has passed."],
        ["cancel(targets, values, calldatas, descriptionHash)", "Same as queue", "Cancel a proposal. Only the proposer (if Pending) or the guardian can cancel."],
    ],
)

doc.add_heading("Governance-Only Methods", level=2)

doc.add_paragraph(
    "These methods can only be called through a successful governance "
    "proposal — they cannot be called directly by any wallet."
)

add_table(
    ["Method", "Parameters", "What It Changes"],
    [
        ["setVotingDelay(newDelay)", "newDelay: uint48", "Changes how many blocks pass before voting begins."],
        ["setVotingPeriod(newPeriod)", "newPeriod: uint32", "Changes how many blocks the voting window stays open."],
        ["setProposalThreshold(newThreshold)", "newThreshold: uint256", "Changes the minimum voting power required to create a proposal."],
        ["updateQuorumNumerator(newNumerator)", "newNumerator: uint256", "Changes the quorum percentage (e.g. 4 → 10 means 10% quorum)."],
        ["setProposalGuardian(newGuardian)", "newGuardian: address", "Changes or removes the proposal guardian. Pass address(0) to remove."],
    ],
)

doc.add_heading("Proposal Guardian", level=2)

doc.add_paragraph(
    "The proposal guardian is an emergency address that can cancel any "
    "non-executed proposal. It acts as a safety brake if a malicious "
    "proposal passes. The guardian cannot create or execute proposals — it "
    "can only cancel them."
)
doc.add_paragraph(
    "Important: The guardian can only be changed or removed through a "
    "governance proposal (setProposalGuardian). If the guardian is "
    "compromised, it could cancel the proposal trying to replace it. "
    "For this reason, always use a multisig as guardian, never a single wallet."
)

# ═══════════════════════════════════════════════════════════════════════════
# 4. WADOOZIE TIMELOCK (TREASURY)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. WadoozieTimelock (Treasury)", level=1)

doc.add_paragraph(
    "The Timelock is the DAO's treasury and execution delay. When a "
    "governance proposal passes, it is queued here and must wait a minimum "
    "delay before it can be executed. This gives the community time to "
    "review approved proposals and exit if they disagree. The Timelock holds "
    "all DAO funds and is the contract that actually executes on-chain "
    "actions."
)

doc.add_heading("Constructor Parameters", level=2)

add_table(
    ["Parameter", "Type", "Description", "Recommended Value"],
    [
        ["minDelay", "uint256", "Minimum seconds a proposal must wait after being queued before it can be executed.", "86,400 (1 day)"],
        ["proposers", "address[]", "Addresses allowed to queue operations. Typically left empty at deploy — the Governor is granted this role afterward.", "[] (empty)"],
        ["executors", "address[]", "Addresses allowed to execute queued operations. Use [address(0)] to allow anyone to execute.", "[address(0)]"],
        ["admin", "address", "Temporary admin who can grant/revoke roles. This role should be renounced after setup.", "deployer (then renounced)"],
    ],
)

doc.add_heading("Read Methods (free, no gas)", level=2)

add_table(
    ["Method", "Returns", "What It Does"],
    [
        ["getMinDelay()", "uint256", "Returns the minimum execution delay in seconds."],
        ["isOperation(id)", "bool", "Returns whether an operation ID exists."],
        ["isOperationPending(id)", "bool", "Returns whether an operation is waiting to be executed."],
        ["isOperationReady(id)", "bool", "Returns whether an operation's delay has passed and it can be executed."],
        ["isOperationDone(id)", "bool", "Returns whether an operation has already been executed."],
        ["getOperationState(id)", "OperationState", "Returns the current state of an operation (Unset, Waiting, Ready, Done)."],
        ["getTimestamp(id)", "uint256", "Returns the timestamp at which an operation becomes executable."],
        ["DEFAULT_ADMIN_ROLE()", "bytes32", "The role hash for the admin role."],
        ["PROPOSER_ROLE()", "bytes32", "The role hash for the proposer role (can queue operations)."],
        ["EXECUTOR_ROLE()", "bytes32", "The role hash for the executor role (can execute operations)."],
        ["CANCELLER_ROLE()", "bytes32", "The role hash for the canceller role (can cancel queued operations)."],
        ["hasRole(role, account)", "bool", "Returns whether an account has a specific role."],
        ["getRoleAdmin(role)", "bytes32", "Returns which role can grant/revoke the given role."],
    ],
)

doc.add_heading("Write Methods (require gas)", level=2)

add_table(
    ["Method", "Parameters", "What It Does"],
    [
        ["schedule(target, value, data, predecessor, salt, delay)", "target: address, value: uint256, data: bytes, predecessor: bytes32, salt: bytes32, delay: uint256", "Queue a single operation for execution after the delay. Requires PROPOSER_ROLE."],
        ["scheduleBatch(targets, values, payloads, predecessor, salt, delay)", "targets: address[], values: uint256[], payloads: bytes[], predecessor: bytes32, salt: bytes32, delay: uint256", "Queue multiple operations as a batch. Requires PROPOSER_ROLE."],
        ["execute(target, value, payload, predecessor, salt)", "target: address, value: uint256, payload: bytes, predecessor: bytes32, salt: bytes32", "Execute a ready operation. Requires EXECUTOR_ROLE (or anyone if role is granted to address(0))."],
        ["executeBatch(targets, values, payloads, predecessor, salt)", "targets: address[], values: uint256[], payloads: bytes[], predecessor: bytes32, salt: bytes32", "Execute a ready batch operation."],
        ["cancel(id)", "id: bytes32", "Cancel a pending operation. Requires CANCELLER_ROLE."],
        ["grantRole(role, account)", "role: bytes32, account: address", "Grant a role to an address. Requires the role's admin role."],
        ["revokeRole(role, account)", "role: bytes32, account: address", "Revoke a role from an address. Requires the role's admin role."],
        ["renounceRole(role, callerAddress)", "role: bytes32, callerAddress: address", "Renounce a role from yourself. Anyone can renounce their own roles."],
        ["updateDelay(newDelay)", "newDelay: uint256", "Change the minimum execution delay. Can only be called by the Timelock itself (i.e., through a governance proposal)."],
    ],
)

doc.add_heading("Roles", level=2)

add_table(
    ["Role", "Who Has It", "What It Can Do"],
    [
        ["DEFAULT_ADMIN_ROLE", "Nobody (renounced after setup)", "Grant and revoke all other roles. Renounced to make the DAO self-governing."],
        ["PROPOSER_ROLE", "Headquarters (Governor)", "Queue approved proposals for execution."],
        ["EXECUTOR_ROLE", "address(0) — anyone", "Execute queued proposals once the delay has passed. Open to anyone by default."],
        ["CANCELLER_ROLE", "Headquarters (Governor)", "Cancel queued proposals that haven't been executed yet."],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════
# 5. DEPLOYMENT STEPS
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Deployment Steps", level=1)

doc.add_paragraph(
    "Follow these steps in order. Each step depends on the previous one."
)

steps = [
    (
        "Deploy the Wadoozie Token",
        "Deploy the Wadoozie contract with the initialHolder address. "
        "The entire 1 billion supply goes to this address."
    ),
    (
        "Deploy the WadoozieTimelock",
        "Deploy with:\n"
        "  - minDelay = 86400 (1 day)\n"
        "  - proposers = [] (empty array)\n"
        "  - executors = [address(0)] (anyone can execute)\n"
        "  - admin = deployer address (temporary)"
    ),
    (
        "Deploy Headquarters",
        "Deploy with the token address, timelock address, and governance "
        "parameters (voting delay, voting period, proposal threshold, quorum "
        "percent, guardian address)."
    ),
    (
        "Grant Timelock Roles to Governor",
        "Call grantRole on the Timelock to give the Governor both "
        "PROPOSER_ROLE and CANCELLER_ROLE. This allows the Governor to "
        "queue and cancel proposals."
    ),
    (
        "Renounce Admin Role",
        "Call renounceRole on the Timelock to remove DEFAULT_ADMIN_ROLE "
        "from the deployer. After this, no human can change roles — only "
        "governance proposals can. This step is critical for decentralization."
    ),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {i}: {title}")
    run.bold = True
    doc.add_paragraph(desc)

doc.add_paragraph()
bold_para("After deployment: ", (
    "The initial token holder should delegate their tokens "
    "(call delegate(ownAddress) on the Wadoozie token) to activate their "
    "voting power, then distribute tokens to community members as planned."
))

# ── Save ──────────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"Created: {OUTPUT}")
