#!/usr/bin/env python3
"""Generate remix_deployment_guide.docx from a structured outline."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "remix_deployment_guide.docx"


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if bold:
        r.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xB0, 0x2A, 0x37)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


# -------------------------------------------------------------------------
doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Title
title = doc.add_heading("Wadoozie — Remix Deployment Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_para(
    doc,
    "Step-by-step procedure for deploying the four Wadoozie contracts "
    "(Wadoozie, WadoozieTreasury, Headquarters, WadoozieDisperser) "
    "from flattened single-file sources using the Remix IDE. "
    "Every step lists the exact field, value, and click required. "
    "Follow the order top to bottom — there are no optional steps "
    "between Section 4 and Section 9.",
)

add_para(doc, " ")

# =========================================================================
add_heading(doc, "0. What you will deploy", level=1)
add_table(
    doc,
    ["#", "Contract", "Constructor args", "Role"],
    [
        ["1", "Wadoozie", "6 addresses", "ERC-20 token (WADZ), ERC20Votes, ERC20Permit"],
        ["2", "WadoozieTreasury", "minDelay, proposers[], executors[], admin", "OZ TimelockController — DAO treasury"],
        ["3", "Headquarters", "token, timelock, 5 numeric params", "OZ Governor — proposes/queues/executes"],
        ["4", "WadoozieDisperser", "(none — no constructor args)", "Stateless ERC-20 batch disperser"],
    ],
)

add_para(doc, " ")
add_para(doc, "Key facts:", bold=True)
add_bullet(doc, "Total supply: 2,000,000,000 WADZ — minted entirely inside the Wadoozie constructor. Deployer holds zero by the end of the deploy tx.")
add_bullet(doc, "999,999,999 WADZ is sent to 0x000…dEaD inside the constructor. totalSupply() stays at 2B; the burn does NOT decrease total supply.")
add_bullet(doc, "LP_WALLET receives 750,000,001 WADZ (the +1 ceremony token).")
add_bullet(doc, "All five recipient addresses on Wadoozie are immutable. They cannot be changed after deploy. Triple-check them before clicking Deploy.")
add_bullet(doc, "WadoozieDisperser is optional — only deploy it if you plan to fan out the Treasury/Publisher/Team allocations or the 576 Signal Fragment recipients in batched txs.")

# =========================================================================
add_heading(doc, "1. Prerequisites", level=1)
add_bullet(doc, "MetaMask (or hardware wallet via MetaMask) connected to the target network — Ethereum mainnet for production, Sepolia for dry runs.")
add_bullet(doc, "Funded deployer EOA — budget ~0.15–0.25 ETH on mainnet at moderate gas. Top up before starting; do not pause mid-flow to refuel.")
add_bullet(doc, "Etherscan API key (for verification step in Section 9). Get one at https://etherscan.io/myapikey.")
add_bullet(doc, "The four flattened source files from this repo: contracts-flat/Wadoozie.sol, WadoozieTreasury.sol, Headquarters.sol, WadoozieDisperser.sol.")
add_bullet(doc, "All five allocation wallet addresses ready (see Section 3). Use multisigs for Treasury, Publisher Rewards, Signal Fragments, and Team Vesting — there is no admin to fix mistakes.")

add_callout(
    doc,
    "Important —",
    "If you have not already, regenerate the flat sources locally with `npm run flatten`. The flat files in contracts-flat/ MUST match the contracts in contracts/ verbatim, otherwise Etherscan verification will fail.",
)

# =========================================================================
add_heading(doc, "2. Open Remix and load the flat sources", level=1)

add_numbered(doc, "Go to https://remix.ethereum.org in a browser.")
add_numbered(doc, "In the left sidebar, click the File Explorer icon (top icon).")
add_numbered(doc, "Right-click on 'contracts/' folder and choose New File. Name it Wadoozie.sol.")
add_numbered(doc, "Open the local file contracts-flat/Wadoozie.sol from this repo, copy its full contents, and paste them into the Remix file you just created. Save with Ctrl+S / Cmd+S.")
add_numbered(doc, "Repeat the previous two steps for each of the other flat files: WadoozieTreasury.sol, Headquarters.sol, WadoozieDisperser.sol. You should end with four files in the Remix file explorer.")

add_callout(
    doc,
    "Tip —",
    "If Remix offers to import from GitHub or a URL, do not. Always paste the LOCAL flat file. The bytes Etherscan sees during verification must match the bytes Remix compiled.",
)

# =========================================================================
add_heading(doc, "3. Decide and record the six addresses", level=1)

add_para(
    doc,
    "These addresses are baked into Wadoozie's constructor and become "
    "immutable after the tx confirms. There is no admin function to "
    "fix any of them later. Write them down somewhere durable BEFORE "
    "deploying.",
)

add_table(
    doc,
    ["Slot", "Description", "Recommended type", "Allocation"],
    [
        ["deployer", "The EOA you will sign the deploy tx with. Briefly holds 2B inside the constructor.", "Hot EOA you control", "0 (drains to zero)"],
        ["LP_WALLET", "Will pair its allocation with ETH on Uniswap immediately post-deploy, then burn the LP receipt.", "Hot EOA / launchpad multisig", "750,000,001 WADZ"],
        ["TREASURY", "DAO-controlled spend wallet.", "Multisig (Gnosis Safe)", "100,000,000 WADZ"],
        ["PUBLISHER_REWARDS", "Pool wallet for Publishers Network creator payouts.", "Multisig", "70,000,000 WADZ"],
        ["SIGNAL_FRAGMENTS", "576-fragment prize pool.", "Multisig", "50,000,000 WADZ"],
        ["TEAM_VESTING", "Team allocation under a 12-month lock.", "Timelock or vesting contract", "30,000,000 WADZ"],
    ],
)

add_callout(
    doc,
    "Burn semantics —",
    "999,999,999 WADZ is transferred to 0x000000000000000000000000000000000000dEaD inside the constructor. "
    "totalSupply() permanently reads 2,000,000,000. Any explorer that displays a 'circulating supply' must "
    "subtract balanceOf(0x...dEaD) manually.",
)

# =========================================================================
add_heading(doc, "4. Configure the Solidity compiler", level=1)

add_numbered(doc, "Click the Solidity Compiler icon in the left sidebar (the second icon — looks like an 'S').")
add_numbered(doc, "In the COMPILER dropdown choose: 0.8.28+commit.7893614a.")
add_numbered(doc, "Click 'Advanced Configurations' to expand it.")
add_numbered(doc, "Set EVM VERSION to: cancun.")
add_numbered(doc, "Tick 'Enable optimization'.")
add_numbered(doc, "Set the optimization runs box to: 200.")
add_numbered(doc, "Leave 'Auto compile' on if you like (or compile manually after each file selection).")
add_numbered(doc, "Click on Wadoozie.sol in the file explorer, then click 'Compile Wadoozie.sol'. Confirm a green check appears.")
add_numbered(doc, "Repeat for WadoozieTreasury.sol, Headquarters.sol, WadoozieDisperser.sol. All four must compile clean — no warnings about license, no errors.")

add_callout(
    doc,
    "Critical —",
    "These exact settings (0.8.28, cancun, optimizer enabled, 200 runs) are what scripts/verify-flat.mjs and the Etherscan UI expect. If you change ANY of them, on-chain bytecode will not match and verification will fail.",
)

# =========================================================================
add_heading(doc, "5. Deploy Wadoozie (the token)", level=1)

add_numbered(doc, "Click the Deploy & Run Transactions icon in the left sidebar (the third icon — looks like an Ethereum logo).")
add_numbered(doc, "ENVIRONMENT dropdown: choose 'Injected Provider - MetaMask'. MetaMask should pop up and ask you to connect; pick the deployer EOA.")
add_numbered(doc, "Confirm the network shown in Remix matches your target network (e.g., 'Custom (1) network' for mainnet, '11155111' for Sepolia).")
add_numbered(doc, "ACCOUNT dropdown: confirm it shows your deployer EOA's address. Confirm a non-zero ETH balance.")
add_numbered(doc, "GAS LIMIT: leave blank (Remix auto-estimates).")
add_numbered(doc, "VALUE: leave at 0 Wei.")
add_numbered(doc, "CONTRACT dropdown: choose 'Wadoozie - contracts/Wadoozie.sol'.")
add_numbered(doc, "Click the orange caret next to the 'Deploy' button to expand the constructor argument fields.")

add_para(doc, "Fill the six constructor fields, in this exact order:", bold=True)

add_table(
    doc,
    ["Field", "Type", "Value to paste"],
    [
        ["deployer_", "address", "<your deployer EOA address — same as the connected MetaMask account>"],
        ["lpWallet_", "address", "<LP_WALLET address>"],
        ["treasury_", "address", "<TREASURY multisig address>"],
        ["publisherRewards_", "address", "<PUBLISHER_REWARDS multisig address>"],
        ["signalFragments_", "address", "<SIGNAL_FRAGMENTS multisig address>"],
        ["teamVesting_", "address", "<TEAM_VESTING address>"],
    ],
)

add_numbered(doc, "Click the orange 'transact' button.")
add_numbered(doc, "MetaMask pops up. Review the network, the contract, and the gas estimate (rough estimate: ~3M gas). If anything looks off, REJECT and re-check the addresses.")
add_numbered(doc, "Click Confirm in MetaMask.")
add_numbered(doc, "Wait for confirmation. Remix shows a green check in the terminal at the bottom and the deployed contract appears under 'Deployed Contracts'.")
add_numbered(doc, "Copy the deployed Wadoozie contract address from Remix (click the copy icon next to it). Save it as TOKEN_ADDRESS in your notes — you will need it in Sections 7 and 9.")

add_callout(
    doc,
    "Sanity-check before continuing —",
    "Expand the deployed Wadoozie under 'Deployed Contracts'. Click totalSupply — it should read 2000000000000000000000000000 (2B × 1e18). Click balanceOf and paste 0x000000000000000000000000000000000000dEaD — it should read 999999999000000000000000000. Click LP_WALLET, TREASURY, PUBLISHER_REWARDS, SIGNAL_FRAGMENTS, TEAM_VESTING — confirm each returns the address you intended.",
)

# =========================================================================
add_heading(doc, "6. Deploy WadoozieTreasury (the Timelock)", level=1)

add_numbered(doc, "Still on the Deploy & Run tab, change CONTRACT dropdown to 'WadoozieTreasury - contracts/WadoozieTreasury.sol'.")
add_numbered(doc, "Click the orange caret next to 'Deploy' to expand the four constructor fields.")

add_para(doc, "Fill the four constructor fields exactly:", bold=True)

add_table(
    doc,
    ["Field", "Type", "Value to paste"],
    [
        ["minDelay", "uint256", "86400"],
        ["proposers", "address[]", "[]"],
        ["executors", "address[]", "[\"0x0000000000000000000000000000000000000000\"]"],
        ["admin", "address", "<your deployer EOA address>"],
    ],
)

add_para(
    doc,
    "Notes on these values:",
    bold=True,
)
add_bullet(doc, "minDelay = 86400 = 1 day in seconds. Queued proposals must wait this long before they can execute.")
add_bullet(doc, "proposers is an empty array — Headquarters will be granted PROPOSER_ROLE in Section 8.")
add_bullet(doc, "executors contains a single address: 0x0000000000000000000000000000000000000000 (zero address). Per OZ semantics this means 'anyone can execute' a queued proposal.")
add_bullet(doc, "admin is your deployer EOA — temporary. You will renounce this role at the end of Section 8.")

add_numbered(doc, "Click the orange 'transact' button.")
add_numbered(doc, "Confirm the tx in MetaMask.")
add_numbered(doc, "After confirmation, copy the WadoozieTreasury contract address. Save it as TIMELOCK_ADDRESS.")

# =========================================================================
add_heading(doc, "7. Deploy Headquarters (the Governor)", level=1)

add_numbered(doc, "Change CONTRACT dropdown to 'Headquarters - contracts/Headquarters.sol'.")
add_numbered(doc, "Click the orange caret next to 'Deploy' to expand the seven constructor fields.")

add_para(doc, "Fill the seven constructor fields exactly:", bold=True)

add_table(
    doc,
    ["Field", "Type", "Value to paste"],
    [
        ["_token", "address (IVotes)", "<TOKEN_ADDRESS from Section 5>"],
        ["_timelock", "address (TimelockController)", "<TIMELOCK_ADDRESS from Section 6>"],
        ["_votingDelay", "uint48", "7200"],
        ["_votingPeriod", "uint32", "50400"],
        ["_proposalThreshold", "uint256", "1000000000000000000000"],
        ["_quorumNumerator", "uint256", "350"],
        ["_voteExtension", "uint48", "7200"],
    ],
)

add_para(doc, "Plain-English meaning of those numbers:", bold=True)
add_bullet(doc, "_votingDelay = 7200 blocks ≈ 1 day on mainnet — delay between proposal creation and start of voting.")
add_bullet(doc, "_votingPeriod = 50400 blocks ≈ 7 days — length of the voting window.")
add_bullet(doc, "_proposalThreshold = 1000 * 1e18 = 1000 WADZ — minimum tokens to create a proposal.")
add_bullet(doc, "_quorumNumerator = 350 — paired with the hard-coded denominator 10000 = 3.5% quorum.")
add_bullet(doc, "_voteExtension = 7200 blocks ≈ 1 day — late-quorum protection (auto-extends voting window if quorum reached late).")

add_numbered(doc, "Click 'transact'.")
add_numbered(doc, "Confirm in MetaMask.")
add_numbered(doc, "After confirmation, copy the Headquarters contract address. Save it as GOVERNOR_ADDRESS.")

add_callout(
    doc,
    "Verify before continuing —",
    "Expand the deployed Headquarters under 'Deployed Contracts'. Click quorumNumerator → 350. Click quorumDenominator → 10000. Click votingDelay → 7200. Click votingPeriod → 50400. If any value differs, you used a wrong constructor argument and must redeploy from this section.",
)

# =========================================================================
add_heading(doc, "8. Wire up roles (PROPOSER, CANCELLER, renounce admin)", level=1)

add_para(
    doc,
    "Right now the timelock has no proposers and the deployer is its only "
    "admin. Three transactions complete the wiring: grant PROPOSER_ROLE to "
    "the governor, grant CANCELLER_ROLE to the governor, then renounce "
    "DEFAULT_ADMIN_ROLE from the deployer.",
)

add_para(doc, "Step 8a — grant PROPOSER_ROLE to Headquarters:", bold=True)
add_numbered(doc, "Under 'Deployed Contracts', expand the WadoozieTreasury instance.")
add_numbered(doc, "Click PROPOSER_ROLE (a blue button — view function). Copy the bytes32 it returns. (For OZ v5 it is 0xb09aa5aeb3702cfd50b6b62bc4532604938f21248a27a1d5ca736082b6819cc1.)")
add_numbered(doc, "Find the grantRole function. role: paste the bytes32 from the previous step. account: paste GOVERNOR_ADDRESS.")
add_numbered(doc, "Click 'transact'. Confirm in MetaMask. Wait for confirmation.")

add_para(doc, "Step 8b — grant CANCELLER_ROLE to Headquarters:", bold=True)
add_numbered(doc, "Click CANCELLER_ROLE (view function) on the same WadoozieTreasury instance. Copy the bytes32. (For OZ v5: 0xfd643c72710c63c0180259aba6b2d05451e3591a24e58b62239378085726f783.)")
add_numbered(doc, "Find grantRole again. role: paste the CANCELLER_ROLE bytes32. account: paste GOVERNOR_ADDRESS.")
add_numbered(doc, "Click 'transact'. Confirm in MetaMask.")

add_para(doc, "Step 8c — renounce DEFAULT_ADMIN_ROLE on the timelock:", bold=True)
add_numbered(doc, "Click DEFAULT_ADMIN_ROLE (view function). It returns 0x0000000000000000000000000000000000000000000000000000000000000000.")
add_numbered(doc, "Find renounceRole. role: paste 0x0000000000000000000000000000000000000000000000000000000000000000. callerConfirmation: paste your deployer EOA address (must equal the caller).")
add_numbered(doc, "Click 'transact'. Confirm in MetaMask. Once this tx confirms, the deployer has no remaining authority over the timelock.")

add_callout(
    doc,
    "Point of no return —",
    "After Step 8c there is no admin on the timelock. Future role changes require a successful governance proposal that targets the timelock itself. Do NOT skip the verifications in Step 8d before walking away.",
)

add_para(doc, "Step 8d — verify the wiring:", bold=True)
add_numbered(doc, "On the timelock, call hasRole with role = PROPOSER_ROLE bytes32, account = GOVERNOR_ADDRESS. Expect true.")
add_numbered(doc, "Call hasRole with role = CANCELLER_ROLE bytes32, account = GOVERNOR_ADDRESS. Expect true.")
add_numbered(doc, "Call hasRole with role = DEFAULT_ADMIN_ROLE bytes32 (all zeros), account = <deployer EOA>. Expect false.")
add_numbered(doc, "If any of those three return the wrong value, STOP. Do not announce the addresses publicly. Investigate before continuing.")

# =========================================================================
add_heading(doc, "9. Verify all contracts on Etherscan", level=1)

add_para(
    doc,
    "Verification publishes the source code so anyone can read it on "
    "Etherscan and so dApps can decode tx data. Repeat the same flow for "
    "each of the deployed contracts.",
)

add_para(doc, "9a — Open the verification page:", bold=True)
add_numbered(doc, "Go to https://etherscan.io/address/<contract_address> (use sepolia.etherscan.io for Sepolia).")
add_numbered(doc, "Click the 'Contract' tab.")
add_numbered(doc, "Click 'Verify and Publish'.")

add_para(doc, "9b — Form values (identical for all four contracts):", bold=True)
add_table(
    doc,
    ["Field", "Value"],
    [
        ["Compiler Type", "Solidity (Single file)"],
        ["Compiler Version", "v0.8.28+commit.7893614a"],
        ["Open Source License Type", "MIT License (MIT)"],
    ],
)
add_numbered(doc, "Click Continue.")

add_para(doc, "9c — Settings on the next page:", bold=True)
add_table(
    doc,
    ["Field", "Value"],
    [
        ["Optimization", "Yes"],
        ["Runs (Optimizer)", "200"],
        ["EVM Version to target", "cancun"],
    ],
)

add_numbered(doc, "Open the matching flat file from contracts-flat/ in a text editor, select all, copy.")
add_numbered(doc, "Paste the entire flat source into the 'Enter the Solidity Contract Code below' box.")
add_numbered(doc, "In 'Constructor Arguments ABI-encoded', paste the hex string for that contract (without the 0x prefix). Generation commands are listed in 9d.")
add_numbered(doc, "Leave 'Contract Library Address' empty (none of these contracts link external libraries).")
add_numbered(doc, "Solve the captcha if shown, then click 'Verify and Publish'.")
add_numbered(doc, "Wait ~10–60 seconds for Etherscan to compile and match. A green check means success. If it fails, see Section 10.")

add_para(doc, "9d — Generate constructor argument hex:", bold=True)
add_para(
    doc,
    "Run this Node.js one-liner in any terminal where ethers is "
    "available (the repo's node_modules will do). It prints all three hex "
    "strings. The disperser has no constructor args, so leave that field "
    "empty when verifying it.",
)
add_code(
    doc,
    "node -e \"\n"
    "  const { AbiCoder } = require('ethers');\n"
    "  const enc = AbiCoder.defaultAbiCoder();\n"
    "  console.log('Wadoozie:', enc.encode(\n"
    "    ['address','address','address','address','address','address'],\n"
    "    ['<DEPLOYER>','<LP_WALLET>','<TREASURY>','<PUBLISHER_REWARDS>',\n"
    "     '<SIGNAL_FRAGMENTS>','<TEAM_VESTING>']\n"
    "  ).slice(2));\n"
    "  console.log('Treasury:', enc.encode(\n"
    "    ['uint256','address[]','address[]','address'],\n"
    "    [86400, [], ['0x0000000000000000000000000000000000000000'], '<DEPLOYER>']\n"
    "  ).slice(2));\n"
    "  console.log('HQ:', enc.encode(\n"
    "    ['address','address','uint48','uint32','uint256','uint256','uint48'],\n"
    "    ['<TOKEN_ADDRESS>','<TIMELOCK_ADDRESS>',7200,50400,\n"
    "     '1000000000000000000000',350,7200]\n"
    "  ).slice(2));\n"
    "\"\n",
)

add_para(doc, "9e — Verify each of the four contracts in turn:", bold=True)
add_bullet(doc, "Wadoozie — paste contracts-flat/Wadoozie.sol; constructor args = the 'Wadoozie:' hex from 9d.")
add_bullet(doc, "WadoozieTreasury — paste contracts-flat/WadoozieTreasury.sol; constructor args = the 'Treasury:' hex from 9d.")
add_bullet(doc, "Headquarters — paste contracts-flat/Headquarters.sol; constructor args = the 'HQ:' hex from 9d.")
add_bullet(doc, "WadoozieDisperser — paste contracts-flat/WadoozieDisperser.sol; constructor args = (leave empty).")

# =========================================================================
add_heading(doc, "10. Troubleshooting verification", level=1)

add_table(
    doc,
    ["Etherscan error", "Likely cause", "Fix"],
    [
        ["Bytecode does not match…", "Compiler / optimizer / EVM settings drifted between Remix and Etherscan.", "Confirm 0.8.28+commit.7893614a, optimizer enabled, runs 200, evm version cancun. Re-paste the flat file, re-run."],
        ["Constructor arguments mismatch", "The hex you pasted does not encode the values used at deploy.", "Regenerate hex from 9d; double-check addresses, especially deployer."],
        ["Already Verified", "Etherscan already accepted an earlier submission.", "Safe to ignore."],
        ["Source file does not contain Contract definition", "The flat file is the wrong file (paste mismatch).", "Make sure you pasted the file whose root contract name matches the deployed contract."],
        ["Library address required", "Should not occur for these contracts.", "If it does, you pasted the wrong source — these contracts link no external libraries."],
    ],
)

# =========================================================================
add_heading(doc, "11. (Optional) Deploying the WadoozieDisperser", level=1)

add_para(
    doc,
    "The disperser is a stateless helper for fanning out ERC-20 transfers "
    "in batches under EIP-7825's per-tx gas cap. Deploy it only if the LP "
    "wallet, Treasury, Publisher Rewards, Team Vesting, or 576 Signal "
    "Fragment recipients need to be funded from the deployer EOA (or any "
    "wallet that holds WADZ).",
)

add_numbered(doc, "On the Deploy & Run tab, change CONTRACT dropdown to 'WadoozieDisperser - contracts/WadoozieDisperser.sol'.")
add_numbered(doc, "There are no constructor arguments — click Deploy directly.")
add_numbered(doc, "Confirm in MetaMask. After confirmation, copy the disperser address as DISPERSER_ADDRESS.")
add_numbered(doc, "From the wallet that holds the tokens you want to disperse, call WADZ.approve(DISPERSER_ADDRESS, totalAmount) where totalAmount is the sum of all transfers in the batch.")
add_numbered(doc, "On the deployed disperser, call disperseTokenSimple(token, recipients, values):")
add_bullet(doc, "token — the WADZ contract address (TOKEN_ADDRESS).")
add_bullet(doc, "recipients — JSON-style array of addresses, e.g. [\"0xabc…\",\"0xdef…\"].")
add_bullet(doc, "values — JSON-style array of uint256 amounts (in 1e18 base units). MUST be the same length as recipients.")

add_callout(
    doc,
    "EIP-7825 cap —",
    "Each disperse call must fit under 16,777,216 gas. With ERC20Votes auto-delegation kicking in for fresh recipients, plan on roughly ~150 transfers per batch as a safe upper bound. Split the 576 Signal Fragment recipients across 4 calls.",
)

# =========================================================================
add_heading(doc, "12. Final post-deploy checklist", level=1)

add_bullet(doc, "All four contracts show a green checkmark on Etherscan.")
add_bullet(doc, "Wadoozie.totalSupply() returns 2000000000000000000000000000 (2,000,000,000 × 1e18).")
add_bullet(doc, "Wadoozie.balanceOf(0x000…dEaD) returns 999999999000000000000000000.")
add_bullet(doc, "Wadoozie.balanceOf(LP_WALLET) returns 750000001000000000000000000.")
add_bullet(doc, "Wadoozie.balanceOf(<deployer EOA>) returns 0.")
add_bullet(doc, "Wadoozie.LP_WALLET / TREASURY / PUBLISHER_REWARDS / SIGNAL_FRAGMENTS / TEAM_VESTING getters return the addresses you intended.")
add_bullet(doc, "WadoozieTreasury.hasRole(PROPOSER_ROLE, GOVERNOR_ADDRESS) returns true.")
add_bullet(doc, "WadoozieTreasury.hasRole(CANCELLER_ROLE, GOVERNOR_ADDRESS) returns true.")
add_bullet(doc, "WadoozieTreasury.hasRole(DEFAULT_ADMIN_ROLE, <deployer>) returns false.")
add_bullet(doc, "Headquarters.token() returns TOKEN_ADDRESS.")
add_bullet(doc, "Headquarters.timelock() returns TIMELOCK_ADDRESS.")
add_bullet(doc, "Headquarters.quorumNumerator() returns 350; quorumDenominator() returns 10000.")
add_bullet(doc, "Deployment record (TOKEN_ADDRESS, TIMELOCK_ADDRESS, GOVERNOR_ADDRESS, optionally DISPERSER_ADDRESS) backed up offline.")
add_bullet(doc, "Public announcement of contract addresses gated until verification is green and the wiring above is confirmed.")

add_para(doc, " ")
add_para(
    doc,
    "There is no upgrade path. The contracts are immutable, the timelock cannot be migrated "
    "(updateTimelock reverts), and there is no admin role left on the treasury. If a bug is "
    "discovered post-deploy, the only mitigation is to deploy v2 and migrate via governance — "
    "token holders vote to send funds out of the existing treasury into a new system.",
)

doc.save(OUT)
print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")
